"""File-content extraction, one function per supported file type.

Each extractor takes a path on disk and returns a JSON-serializable dict.
Called from the Celery task in app.workers.tasks — kept import-light and
side-effect-free (no DB access) so it's independently unit-testable.
"""
from pathlib import Path
from typing import Any

import pandas as pd

from app.core.logging import logger
from app.models.upload import FileType

# Cap how much we inline in the extracted JSON for large tabular files —
# the full data stays queryable from the original upload; this JSON is a
# preview + structure summary for the ML/NLP phases to build on.
MAX_PREVIEW_ROWS = 200


def extract_csv(path: Path) -> dict[str, Any]:
    df = pd.read_csv(path)
    return _dataframe_summary(df)


def extract_xlsx(path: Path) -> dict[str, Any]:
    sheets = pd.read_excel(path, sheet_name=None)
    return {
        "sheet_names": list(sheets.keys()),
        "sheets": {name: _dataframe_summary(df) for name, df in sheets.items()},
    }


def extract_pdf(path: Path) -> dict[str, Any]:
    import pdfplumber

    pages: list[dict[str, Any]] = []
    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            pages.append(
                {
                    "page_number": page_number,
                    "text": page.extract_text() or "",
                    "table_count": len(page.extract_tables()),
                }
            )
    full_text = "\n".join(p["text"] for p in pages)
    return {"page_count": len(pages), "pages": pages, "full_text": full_text}


def extract_docx(path: Path) -> dict[str, Any]:
    import docx

    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    tables = [
        [[cell.text for cell in row.cells] for row in table.rows] for table in document.tables
    ]
    return {
        "paragraph_count": len(paragraphs),
        "paragraphs": paragraphs,
        "table_count": len(tables),
        "tables": tables,
        "full_text": "\n".join(paragraphs),
    }


def extract_image(path: Path) -> dict[str, Any]:
    """OCR placeholder.

    No OCR engine is wired up yet (Tesseract/cloud OCR lands in a later
    phase). For now this records image metadata and an explicit flag so
    downstream consumers know no text was extracted and shouldn't treat
    an empty result as "no text in the image".
    """
    from PIL import Image

    with Image.open(path) as img:
        width, height = img.size
        image_format = img.format

    logger.info(f"OCR not yet implemented — recorded metadata only for {path.name}")
    return {
        "ocr_placeholder": True,
        "width": width,
        "height": height,
        "format": image_format,
        "text": None,
    }


_EXTRACTORS = {
    FileType.CSV: extract_csv,
    FileType.XLSX: extract_xlsx,
    FileType.PDF: extract_pdf,
    FileType.DOCX: extract_docx,
    FileType.PNG: extract_image,
    FileType.JPG: extract_image,
}


def extract_content(file_type: FileType, path: Path) -> dict[str, Any]:
    """Dispatches to the right extractor. Raises whatever the extractor raises —
    the Celery task is responsible for catching it and marking the upload FAILED.
    """
    extractor = _EXTRACTORS[file_type]
    return extractor(path)


def _dataframe_summary(df: pd.DataFrame) -> dict[str, Any]:
    return {
        "row_count": int(len(df)),
        "column_count": int(len(df.columns)),
        "columns": [str(c) for c in df.columns],
        "dtypes": {str(c): str(dtype) for c, dtype in df.dtypes.items()},
        "preview": df.head(MAX_PREVIEW_ROWS).to_dict(orient="records"),
    }
